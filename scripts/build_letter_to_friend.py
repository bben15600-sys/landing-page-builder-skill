#!/usr/bin/env python3
"""Short, casual 'what I'm building' doc — for sending to a friend.

Not marketing, not engineering — just honest explanation, 2-3 pages.
Hebrew, friendly tone, one concrete example.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x0B, 0x5E, 0xD7)
MUTED = RGBColor(0x6B, 0x72, 0x80)
SOFT = RGBColor(0x4B, 0x55, 0x63)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def para(doc, text, size=12, bold=False, italic=False, color=None, space_after=8, rtl=True):
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    return p


def section_title(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    return p


def soft_box(doc, text, italic=True):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, 'F3F6FB')
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CBD5E1')
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = SOFT
    run.font.name = 'Calibri'
    p.paragraph_format.line_spacing = 1.5


def example_box(doc, prompt, result):
    table = doc.add_table(rows=2, cols=1)

    c1 = table.rows[0].cells[0]
    set_cell_bg(c1, '1A1A2E')
    p1 = c1.paragraphs[0]
    set_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0 = p1.add_run('אני: ')
    r0.font.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    r0.font.name = 'Calibri'
    r1 = p1.add_run(prompt)
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.name = 'Calibri'

    c2 = table.rows[1].cells[0]
    set_cell_bg(c2, 'F3F6FB')
    p2 = c2.paragraphs[0]
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2_label = p2.add_run('הכלי: ')
    r2_label.font.bold = True
    r2_label.font.size = Pt(11)
    r2_label.font.color.rgb = ACCENT
    r2_label.font.name = 'Calibri'
    r2 = p2.add_run(result)
    r2.font.size = Pt(11.5)
    r2.font.color.rgb = INK
    r2.font.name = 'Calibri'
    p2.paragraph_format.line_spacing = 1.5


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # Title (no full cover, just nice header)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('היי, בא לי לספר לך על משהו שאני בונה')
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = INK
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run('פרויקט אישי. עוד לא התחלתי לקוד, אבל יש לי תכנון מוצק. רוצה שמועה?')
    r2.font.size = Pt(12)
    r2.font.italic = True
    r2.font.color.rgb = MUTED
    r2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(20)

    # --- Section: what bothers me ---
    section_title(doc, 'מאיפה זה בא')
    para(doc,
         'בשנה האחרונה אני משתמש בכלי AI לקוד — Cursor, Claude Code, Lovable, Bolt. '
         'כל אחד מהם עושה משהו טוב, אבל לכל אחד יש את החיסרון שלו:')
    para(doc, '•  Lovable ו-Bolt — יפים מאוד, אבל נועלים אותך על React + Supabase. לא עוזר אם בא לי לבנות משהו ב-Python או mobile.', size=11)
    para(doc, '•  Claude Code — חזק מאוד, אבל $100 לחודש עם limit של ~225 הודעות כל 5 שעות. נגמר מהר.', size=11)
    para(doc, '•  Cursor — נחמד, אבל הוא IDE, לא סוכן עצמאי שמייצר פרויקטים שלמים.', size=11)
    para(doc, '•  כולם תלויים בספק AI אחד. אם Anthropic מעלה מחיר או Claude יורד — אתה תקוע.', size=11, space_after=14)

    soft_box(doc,
        'בקיצור, הרגשתי שאני שוכר את הכלים שלי במקום להחזיק בהם. אז חשבתי — '
        'מה אם אבנה לעצמי את הכלי שאני באמת רוצה?')

    # --- Section: what it is ---
    section_title(doc, 'מה זה הולך להיות')
    para(doc,
         'CodeForge — סוכן AI אישי שרץ בדפדפן. כותב לו בעברית מה אתה רוצה, '
         'הוא בונה את זה — אתר, אפליקציה, סקריפט, מה שלא יהיה.')

    para(doc, 'מה שמייחד אותו (לפחות בעיני שלי):', bold=True, space_after=4)
    para(doc, '•  לא נעול על framework. בונה ב-Next.js, Python, React Native, Go — מה שצריך.', size=11)
    para(doc, '•  מתחלף בין מודלים. Claude לתכנון, Haiku לעריכות זריזות, Gemini ל-codebase ענק, Qwen מקומי לפרטיות.', size=11)
    para(doc, '•  Preview חי תוך כדי בנייה — הקוד רץ בדפדפן עצמו, אני רואה את האפליקציה תוך שניות.', size=11)
    para(doc, '•  מחובר ל-GitHub ו-Vercel שלי — push, deploy, הכל מהצ\'אט.', size=11)
    para(doc, '•  זוכר אותי. אחרי שבוע אני חוזר ל-session, רואה את כל ההיסטוריה, ממשיך מהמקום שעצרתי.', size=11, space_after=14)

    # --- Section: example ---
    section_title(doc, 'נראה איך זה הולך להיות בפועל')
    para(doc, 'בוא נגיד שאני רוצה אפליקציה לעקוב אחרי הרגלים. במקום לפתוח VS Code ולכתוב boilerplate חצי שעה, אני פשוט אומר:', space_after=8)

    example_box(doc,
        'תבנה לי אפליקציה פשוטה לעקוב אחרי הרגלים. רוצה Next.js, dark mode, '
        'מאגר ב-Postgres, להתחבר עם GitHub. ושמעבירים את הכפתור — מתפוצצים confetti.',
        '✓ יצרתי Next.js 15 project\n'
        '✓ הוספתי Tailwind + shadcn/ui עם dark mode toggle\n'
        '✓ הגדרתי Drizzle schema (habits, completions)\n'
        '✓ Postgres רץ ב-Neon — חיברתי את ה-connection string\n'
        '✓ NextAuth עם GitHub OAuth מוגדר\n'
        '✓ Confetti.js עם trigger על completion\n'
        '\n'
        'הכל רץ ב-localhost כרגע. רוצה לראות?  ▶ פתח preview\n'
        'רוצה להעלות ל-Vercel? אני מטפל. (15 שניות)')

    para(doc, ' ', size=4)
    para(doc, 'תוך 5-10 דקות — אפליקציה רצה, יש URL חי, יש repo פרטי. במקום שעות של setup.', italic=True, color=MUTED)

    doc.add_page_break()

    # --- Section: how it works (simple) ---
    section_title(doc, 'איך זה עובד מבפנים (בלי להעמיק יותר מדי)')
    para(doc,
         'אני בונה את זה כ-web app ב-Next.js שעובד ככה:')
    para(doc, '1.  אני כותב הודעה. היא נשלחת ל-API.', size=11)
    para(doc, '2.  ה-API מדבר עם Claude (או GPT, או Gemini — תלוי במשימה) ומקבל בחזרה לא רק טקסט אלא גם "פקודות לכלים" — תצור קובץ, הרץ פקודה, פתח דפדפן.', size=11)
    para(doc, '3.  הכלים האלה רצים בתוך sandbox שחי בדפדפן עצמו (טכנולוגיה שנקראת WebContainers). זה אומר אפס עלות שרת, ובידוד מלא.', size=11)
    para(doc, '4.  התוצאה חוזרת לסוכן, הוא מסיק מסקנה, ממשיך — עד שהמשימה הושלמה.', size=11)
    para(doc, '5.  כל הצ\'אט והקבצים נשמרים ב-Postgres, אז אפשר לחזור אחרי שבוע ולהמשיך.', size=11, space_after=14)

    soft_box(doc,
        'הטריק החשוב: המערכת מדברת עם הספקים האלה דרך adapter — אז ביום שבו '
        'Claude נופל או מתייקר, אני מחליף את ה-router שלי ל-Gemini, '
        'והמערכת ממשיכה לעבוד בלי שאני אצטרך לשנות כלום מצד המשתמש.')

    # --- Section: budget reality ---
    section_title(doc, 'תכלס — כמה זה יעלה לי')
    para(doc,
         'זה לא startup. זה כלי בשבילי. אז התקצוב שלי הוא <$150/חודש, וזה ריאלי כי:')
    para(doc, '•  Vercel, GitHub, Neon — חינם בtier שלי.', size=11)
    para(doc, '•  WebContainers — חינם, רץ ב-browser, אפס compute server.', size=11)
    para(doc, '•  ה-API של Anthropic — עם prompt caching (90% הנחה על cache reads) זה ~$30-50/חודש בשימוש כבד.', size=11)
    para(doc, '•  לעומת Claude Code Max ב-$100 עם limits, אני מקבל יותר חופש בפחות כסף.', size=11, space_after=14)

    # --- Section: where I am ---
    section_title(doc, 'איפה אני עכשיו')
    para(doc, 'יש לי תכנית הנדסית מפורטת ל-12 שבועות:', space_after=4)
    para(doc, '•  שבועות 1-3: foundation, chat, וכלים בסיסיים', size=11)
    para(doc, '•  שבועות 4-7: sandbox, editor, preview, GitHub, deploy', size=11)
    para(doc, '•  שבועות 8-10: multi-provider, MCP, cost tracking', size=11)
    para(doc, '•  שבועות 11-12: memory, polish, dogfooding', size=11, space_after=14)
    para(doc,
         'אחרי 12 שבועות — אני מתחיל להשתמש בכלי במקום ב-Claude Code, '
         'ובמשך חודש לא מוסיף פיצ\'רים, רק משתמש. ככה אני יודע מה באמת חסר ומה אני סתם חשבתי שיהיה כיף.')

    # --- Section: why I'm telling you ---
    section_title(doc, 'למה אני מספר לך')
    para(doc,
         'כי אני מתחיל פרויקט ארוך, ומנסיון פרויקטים אישיים יוצאים יותר טוב כשמישהו אחד יודע ושואל אותך פעם בשבועיים "אז מה קורה עם הכלי שלך?"')
    para(doc,
         'אז זהו, רציתי שתדע. אם בא לך לעקוב, או לחשוב איתי על אדריכלות, או פשוט לשלוח לי memes על Cursor — אני פה.')

    para(doc, ' ', size=4)
    p_final = doc.add_paragraph()
    set_rtl(p_final)
    p_final.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf = p_final.add_run('🛠️')
    rf.font.size = Pt(20)
    p_final.paragraph_format.space_after = Pt(0)

    out = '/home/user/landing-page-builder-skill/CodeForge-Letter-To-Friend.docx'
    doc.save(out)
    print(f'✅ Saved: {out}')


if __name__ == '__main__':
    build()
