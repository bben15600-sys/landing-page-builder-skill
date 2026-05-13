#!/usr/bin/env python3
"""Generate a beautifully formatted Word document for CodeForge."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x0B, 0x5E, 0xD7)
ACCENT = RGBColor(0xE8, 0x8B, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x72, 0x80)
SUCCESS = RGBColor(0x10, 0xA8, 0x5A)
HEADER_BG = "0B5ED7"
TABLE_HEADER_BG = "1A1A2E"
ALT_ROW_BG = "F3F6FB"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0B5ED7')
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text, level=1, color=None, size=None):
    color = color or PRIMARY
    sizes = {0: 32, 1: 22, 2: 16, 3: 13}
    size = size or sizes.get(level, 14)
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(doc, text, bold=False, color=None, size=11, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.runs[0] if p.runs else p.add_run('')
    p.runs[0].text = ''
    new_run = p.add_run(text)
    new_run.font.size = Pt(11)
    new_run.font.name = 'Calibri'
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], TABLE_HEADER_BG)
        cell_p = hdr_cells[i].paragraphs[0]
        set_rtl(cell_p)
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        if r_idx % 2 == 0:
            for c in cells:
                set_cell_bg(c, ALT_ROW_BG)
        for c_idx, val in enumerate(row):
            cell_p = cells[c_idx].paragraphs[0]
            set_rtl(cell_p)
            cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell_p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for r in table.rows:
                r.cells[i].width = Inches(w)
    return table


def stats_grid(doc, items):
    """3-column colored stats grid (number + label per cell)."""
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)

    for i, (num, label) in enumerate(items):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, ALT_ROW_BG)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '6')
            b.set(qn('w:color'), '0B5ED7')
            tc_borders.append(b)
        tc_pr.append(tc_borders)

        p1 = cell.paragraphs[0]
        set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(num)
        r1.font.size = Pt(24)
        r1.font.bold = True
        r1.font.color.rgb = PRIMARY
        r1.font.name = 'Calibri'

        p2 = cell.add_paragraph()
        set_rtl(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = MUTED
        r2.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def stack_diagram(doc, rows):
    """5-row colored stack with descending shade — visual architecture layers."""
    colors = ['0B5ED7', '1E40AF', '5B21B6', '7C2D12', '1A1A2E']
    table = doc.add_table(rows=len(rows), cols=1)
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)

    for i, (label, desc) in enumerate(rows):
        cell = table.rows[i].cells[0]
        set_cell_bg(cell, colors[i % len(colors)])

        p1 = cell.paragraphs[0]
        set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(label)
        r1.font.size = Pt(13)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r1.font.name = 'Calibri'

        p2 = cell.add_paragraph()
        set_rtl(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(desc)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r2.font.name = 'Calibri'
    return table


def cover_page(doc):
    """Full-width colored cover banner with title + tagline + quote."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, '0B5ED7')

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(60)
    r1 = p1.add_run('CodeForge')
    r1.font.size = Pt(72)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p2)
    r2 = p2.add_run('מנוע ה-AI האוניברסלי לבניית כל דבר')
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2.font.name = 'Calibri'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p3)
    p3.paragraph_format.space_before = Pt(20)
    r3 = p3.add_run('מערכת אחת. כל מסגרת פיתוח. כל פלטפורמה. כל פרויקט.\nמ-prompt בעברית לאפליקציה חיה בייצור — תוך דקות.')
    r3.font.size = Pt(13)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0xE7, 0xF1, 0xFE)
    r3.font.name = 'Calibri'

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p4)
    p4.paragraph_format.space_before = Pt(30)
    p4.paragraph_format.space_after = Pt(60)
    r4 = p4.add_run('מסמך תקציר · מבוסס על 15 מחקרים עצמאיים מעמיקים')
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0xE7, 0xF1, 0xFE)
    r4.font.name = 'Calibri'
    return table


def callout_box(doc, title, text, bg="E7F1FE", border="0B5ED7"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), border)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    p1 = cell.paragraphs[0]
    set_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = PRIMARY
    r1.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'
    return table


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # COVER — full-width colored banner
    cover_page(doc)
    doc.add_page_break()

    # SECTION 1 — מה זה
    heading(doc, '🎯 מה זה CodeForge?', level=1)
    para(doc,
         'CodeForge היא מערכת AI שבונה תוכנה כמו מהנדס בכיר — מקבלת תיאור בשפה טבעית '
         'ומייצרת פרויקטים מלאים: אתרים, אפליקציות web ו-mobile, backends, integrations, ו-data pipelines. '
         'עם חיבור ל-MCP, ניתוב חכם בין מודלי AI מרובים, ו-deploy לייצור בלחיצת כפתור.')
    para(doc,
         'זו לא רק "עוד צ\'אט AI לקוד" — זו תשתית שמשלבת את הטוב מ-Claude Code, v0, Bolt ו-Devin, '
         'ושמה את הבעלות, הגמישות והאוטונומיה בידיים שלך.')

    doc.add_paragraph()
    stats_grid(doc, [
        ('17', 'יכולות זהות ל-Claude Code'),
        ('10+', 'אינטגרציות MCP מ-day one'),
        ('60-80%', 'חיסכון בעלות API'),
    ])
    doc.add_paragraph()

    # SECTION 2 — 5 יתרונות
    heading(doc, '✨ במה היא מיוחדת — 5 יתרונות תחרותיים', level=1)

    advantages = [
        ('🌐 לא נעולה למסגרת אחת',
         'v0, Bolt, Lovable — כולם תקועים על React + Next.js + Supabase. CodeForge בונה בכל דבר: '
         'Next.js, SvelteKit, Vue, Django, Rails, Go, FastAPI, React Native, Expo, Electron, Unity, CLIs, ML pipelines.'),
        ('🔌 MCP-Native מהיום הראשון',
         'אינטגרציה מובנית ל-10 שירותים מובילים: GitHub, Vercel, Cloudflare, Postgres, Supabase, Stripe, Figma, Playwright, Sentry, Filesystem+Git. '
         'OAuth אוטומטי, hash-pinning להגנה, audit log מלא.'),
        ('🧠 מנוע Multi-Provider חכם',
         'לא תלוי במודל אחד. ניתוב אוטומטי בין Claude Opus/Sonnet/Haiku + GPT-5.5 + Gemini 3.1 + Qwen3-Coder לפי המשימה. '
         'חיסכון של 60-80% בעלות API לעומת שימוש במודל קצה אחד.'),
        ('🪞 חשיפה כ-MCP Server (Distribution Play)',
         'המערכת לא רק צורכת MCP — היא מספקת אותו. משתמשי Cursor, Claude Code, ChatGPT ו-Zed יכולים לקרוא ל-CodeForge מתוך ה-IDE שלהם. '
         'הופך אותנו לתשתית של כל מי שכבר עובד עם AI ב-IDE.'),
        ('🔬 Verification עם דפדפן אמיתי',
         'v0/Bolt עוצרים ב-typecheck. CodeForge פותח דפדפן, לוחץ על כפתורים ובודק שזה עובד. '
         'תבנית Playwright Test Agents (אנת\'רופיק, Oct 2025): plan → generate → execute → heal.'),
    ]
    for title, txt in advantages:
        callout_box(doc, title, txt)
        doc.add_paragraph()

    # SECTION 3 — Multi-Provider Routing Table
    heading(doc, '🧠 ניתוב חכם — איזה מודל למה?', level=2)
    make_table(doc,
        ['משימה', 'מודל ראשי', 'סיבה'],
        [
            ['תכנון ארכיטקטוני', 'Claude Opus 4.7', 'חשיבה עמוקה'],
            ['כתיבת קוד יומיומית', 'Claude Sonnet 4.6', 'מאוזן, חסכוני'],
            ['עריכות פשוטות', 'Claude Haiku 4.5', 'מהיר, זול פי 10'],
            ['Codebase ענק (>500K שורות)', 'Gemini 3.1 Pro', 'חלון 1M tokens'],
            ['Terminal-heavy', 'GPT-5.5', 'הכי טוב ב-shell'],
            ['Privacy / On-prem', 'Qwen3-Coder 480B', 'open weights'],
        ],
        col_widths=[1.8, 2.0, 2.3])

    doc.add_page_break()

    # SECTION 4 — Capabilities vs Claude Code
    heading(doc, '🛠️ מה היא יודעת לעשות — מול Claude Code', level=1)
    para(doc, 'כל מה ש-Claude Code יודע — CodeForge יודעת. ועוד.', bold=True, color=SUCCESS, size=13)
    doc.add_paragraph()

    make_table(doc,
        ['יכולת', 'CodeForge', 'Claude Code'],
        [
            ['קריאת codebases עד 1M tokens', '✅', '✅'],
            ['עריכות multi-file עם diffs מדויקים', '✅', '✅'],
            ['הרצת Bash בסביבה מבוקרת', '✅', '✅'],
            ['חיפוש Web ושליפת דוקומנטציה', '✅', '✅'],
            ['MCP integrations (10+ שירותים)', '✅', '✅'],
            ['Multi-agent fan-out לחקירה במקביל', '✅', '✅'],
            ['Plan Mode עם אישור משתמש', '✅', '✅'],
            ['זיכרון ארוך-טווח (CLAUDE.md / SKILL.md)', '✅', '✅'],
            ['Skills library מתחזק את עצמו', '✅', '✅'],
            ['Hooks ואוטומציות מותאמות', '✅', '✅'],
            ['בדיקת UI בדפדפן אמיתי בלולאה', '✅', 'חלקי'],
            ['Live preview גרפי תוך כדי בנייה', '✅', '❌'],
            ['Multi-provider routing (לא רק Claude)', '✅', '❌'],
            ['Deploy מובנה ל-Vercel/CF/Netlify', '✅', '❌'],
            ['ממשק Web ידידותי ללא-מפתחים', '✅', '❌'],
            ['ניהול ארגונים, צוותים, billing', '✅', '❌'],
            ['BYOK (משתמשים מביאים API keys)', '✅', 'חלקי'],
            ['חשיפה כ-MCP server לכלים אחרים', '✅', '❌'],
        ],
        col_widths=[3.5, 1.3, 1.3])

    doc.add_paragraph()

    heading(doc, '🌐 מה אפשר לבנות עם זה', level=2)
    use_cases = [
        '🛍️ אתרי מסחר — Shopify-like, Stripe מובנה, payment flows',
        '📱 אפליקציות SaaS — auth, billing, dashboards, multi-tenancy',
        '📊 כלי פנים-ארגוניים — admin panels, CRMs, BI dashboards',
        '🤖 AI apps — chatbots, agents, RAG over data',
        '🎮 משחקים ב-browser — Phaser, Three.js, WebGPU',
        '📲 אפליקציות mobile — Expo + React Native',
        '🔌 APIs ו-backends — REST, GraphQL, tRPC, gRPC',
        '🛠️ CLI tools ו-scripts — Node, Python, Go',
        '🧬 Data pipelines — ETL, scrapers, analytics',
    ]
    for uc in use_cases:
        bullet(doc, uc)

    doc.add_page_break()

    # SECTION 5 — COST
    heading(doc, '💰 השוואת עלויות מול Claude Code', level=1)
    callout_box(doc,
        '❓ השאלה הכי חשובה — האם זה יעלה לי יותר?',
        'תשובה קצרה: לא. ובטווח הארוך — הרבה פחות. אתה ה-Engine. אתה קובע איך לשלם.',
        bg="FFF5E0", border="E88B00")
    doc.add_paragraph()

    heading(doc, '📊 שלוש דרכים להפעיל את CodeForge', level=2)
    make_table(doc,
        ['מודל', 'עלות חודשית', 'יתרון'],
        [
            ['BYOK (Anthropic API ישיר)', 'רק token usage', 'אותו מחיר של Claude Code, בלי מגבלות messages'],
            ['Multi-provider routing', '60-80% פחות מ-Anthropic ישיר', 'ניתוב חכם למודל הזול ביותר שמספיק'],
            ['Open weights (Qwen/DeepSeek)', '~$0 (self-hosted)', 'עלות חומרה בלבד, פרטיות מלאה'],
        ],
        col_widths=[2.4, 2.0, 2.6])

    doc.add_paragraph()
    heading(doc, '🎯 דוגמה מספרית', level=2)

    make_table(doc,
        ['פרמטר', 'Claude Code Max 5x', 'CodeForge עם BYOK'],
        [
            ['מחיר חודשי', '$100 קבוע', '$30-50 (לפי שימוש)'],
            ['מגבלת הודעות', '~225 / 5 שעות', 'אין מגבלה'],
            ['מספר ספקי AI', '1 (Anthropic)', '6 ספקים'],
            ['חיסכון יחסי', '—', '$50-70/חודש'],
        ],
        col_widths=[2.0, 2.3, 2.7])

    doc.add_paragraph()
    callout_box(doc,
        '🚀 הבונוס של Prompt Caching',
        'Anthropic נותן 90% הנחה על cache reads. עבור system prompt של 50K tokens שחוזר על עצמו — '
        'חיסכון של 70-90% מעלות ה-input. CodeForge מנצל את זה כברירת מחדל.',
        bg="E7FBE7", border="10A85A")

    doc.add_page_break()

    # SECTION 6 — Architecture
    heading(doc, '🏛️ ארכיטקטורת המערכת', level=1)
    para(doc, 'חמש שכבות מובנות, כל אחת ניתנת להחלפה:', bold=True)
    doc.add_paragraph()

    stack_diagram(doc, [
        ('1. שכבת ממשק (Frontend)',
         'Next.js 15 + AI SDK v5 · Chat · Monaco Editor · Live Preview · File Tree'),
        ('2. שכבת Agent',
         'Single-agent harness + sub-agents · כלים: Read · Edit · Write · Bash · Grep · Web · Browser'),
        ('3. שכבת אינטליגנציה',
         'Multi-Provider Routing · MCP Gateway (10+ servers) · Sandbox (Firecracker microVMs)'),
        ('4. שכבת Storage & Services',
         'Neon Postgres + pgvector · Cloudflare R2 · WorkOS Auth · Stripe · Inngest'),
        ('5. שכבת Deploy',
         'Vercel · Cloudflare Pages · Netlify · Custom infrastructure'),
    ])

    doc.add_paragraph()
    heading(doc, '🧰 Tech Stack מומלץ', level=2)
    make_table(doc,
        ['שכבה', 'טכנולוגיה', 'סיבה'],
        [
            ['Frontend', 'Next.js 15 + TypeScript + Tailwind', 'תקן התעשייה לסוכני AI'],
            ['Streaming', 'Vercel AI SDK v5', 'typed end-to-end, resumable'],
            ['Editor', 'Monaco (desktop) + CodeMirror 6 (mobile)', 'VS Code DX'],
            ['Database', 'Neon Postgres + pgvector', 'branching ל-preview environments'],
            ['ORM', 'Drizzle', 'edge-compatible, מהיר'],
            ['Auth', 'WorkOS AuthKit', 'SSO חינמי עד 1M MAU'],
            ['Billing', 'Stripe + Paddle (MoR)', 'תמיכה ב-VAT ישראלי'],
            ['Sandbox', 'WebContainers + E2B', 'preview מיידי + heavy ops'],
            ['Observability', 'Langfuse + OpenLLMetry', 'OTel-native'],
            ['Jobs', 'Inngest', 'durable execution לסוכנים ארוכים'],
        ],
        col_widths=[1.5, 2.6, 2.7])

    doc.add_page_break()

    # SECTION 7 — Security
    heading(doc, '🔐 אבטחה ברמת Enterprise', level=1)
    para(doc, 'שבע שכבות הגנה מובנות מהיום הראשון:', bold=True)

    sec_items = [
        ('🛡️ Credential Proxy', 'המודל אף פעם לא רואה סודות אמיתיים — placeholders בלבד, החלפה במוצא'),
        ('🏰 Sandbox-per-session', 'Firecracker microVMs עם בידוד קרנל מלא לכל הרצה'),
        ('🚫 Deny Rules', '`rm -rf`, `git push --force`, `DROP TABLE` חסומים בקוד'),
        ('🔗 MCP Hash-Pinning', 'הגנה מפני התקפת MCPoison (CVE-2025-54136)'),
        ('🤖 Pre-execution Classifier', 'Lakera + Llama Prompt Guard 2 לסינון פקודות הרסניות'),
        ('📋 Audit Log מלא', 'ClickHouse עם Merkle chain חתום לכל פעולה'),
        ('✅ SOC 2 Type 2 Track', 'מהיום הראשון עם Vanta, מוכן ל-enterprise תוך 9-12 חודשים'),
    ]
    for t, txt in sec_items:
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p.add_run(t + ' — ')
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = ACCENT
        r1.font.name = 'Calibri'
        r2 = p.add_run(txt)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'

    doc.add_paragraph()

    # SECTION 8 — Business Model
    heading(doc, '📈 מודל עסקי', level=1)
    make_table(doc,
        ['', 'Free', 'Pro', 'Team', 'Enterprise'],
        [
            ['מחיר', '$0', '$25/חודש', '$35/seat', 'Custom'],
            ['Credits/חודש', '30', '150', '200/seat', '∞'],
            ['פרויקטים פרטיים', '❌', '✅', '✅', '✅'],
            ['Custom domain', '—', '1', '5', '∞'],
            ['SSO', '—', '—', 'Google', 'SAML/SCIM'],
            ['BYOK', '—', '—', '—', '✅'],
            ['SLA', '—', '—', '—', '99.9%'],
        ],
        col_widths=[1.5, 1.0, 1.3, 1.3, 1.4])

    doc.add_paragraph()
    callout_box(doc,
        '💎 כלכלת היחידה',
        '3x markup על עלות API → 65-70% gross margin. עם prompt caching אגרסיבי, '
        'יעד $25 ARPU עם $7-9 עלות = $16-18 רווח גולמי לחודש לכל משתמש Pro.',
        bg="E7F1FE", border="0B5ED7")

    doc.add_page_break()

    # SECTION 9 — Timeline
    heading(doc, '🚀 לוח זמנים — מ-0 ל-MVP בשישה חודשים', level=1)
    make_table(doc,
        ['חודש', 'שלב', 'תוצר'],
        [
            ['1', 'Foundation', 'Next.js + Postgres + Auth + Stripe'],
            ['2', 'Agent Core', 'Streaming loop + 7 כלים + caching'],
            ['3', 'UI/UX', 'Chat + Monaco + file tree + iframe preview'],
            ['4', 'Execution', 'WebContainers + E2B + Playwright verify'],
            ['5', 'Hardening', 'Multi-provider + Observability + Security'],
            ['6', 'Differentiation', 'MCP gateway + Skills library + A/B + SOC 2'],
        ],
        col_widths=[0.8, 1.6, 4.0])

    doc.add_paragraph()

    # SECTION 10 — Why Now
    heading(doc, '🎯 למה זה הזמן הנכון', level=1)
    why_now = [
        '📊 SWE-bench Verified הגיע ל-82% ב-2026 — סוכנים אוטונומיים על משימות אמיתיות',
        '💵 Lovable עשתה $100M ARR ב-8 חודשים, Cursor חצתה $500M ARR',
        '🌍 MCP הפך לתקן — Anthropic, OpenAI, Cursor, Zed, Windsurf, ChatGPT',
        '🔑 תעלת אמיתית נסגרת מהר — מי שלא משחק עכשיו לא ישחק בכלל',
        '🇮🇱 שוק ישראל חסר פתרון מקומי עם תמיכה בעברית ו-RTL',
    ]
    for w in why_now:
        bullet(doc, w)

    doc.add_paragraph()

    # SECTION 11 — Bottom Line
    heading(doc, '💡 שורה תחתונה', level=1)
    para(doc,
        'CodeForge הופכת אותך מ-משתמש של Claude Code ל-בעלים של מנוע AI עצמאי.',
        bold=True, color=PRIMARY, size=14)

    doc.add_paragraph()

    bottom = [
        '✅ כל היכולות של Claude Code — וגם יותר',
        '✅ עלות נמוכה יותר תוך שליטה מלאה',
        '✅ ממשק Web יפה לא-מפתחים, API למפתחים',
        '✅ אינטגרציה עם 10+ שירותים דרך MCP',
        '✅ אסטרטגיית distribution דרך חשיפת המערכת כ-MCP server',
        '✅ אבטחה ו-compliance ברמת enterprise',
    ]
    for b in bottom:
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(b)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = SUCCESS
        r.font.name = 'Calibri'

    doc.add_paragraph()
    callout_box(doc,
        '🚀 צעד אחד מהפעלה',
        'זה לא רק "Claude Code שלי". זו הפלטפורמה שעליה אבנה את כל מה שעוד יבוא. '
        'אישור החזון → Phase 1 (Foundation) → MVP חי תוך 90 יום.',
        bg="FFF5E0", border="E88B00")

    doc.add_paragraph()
    add_horizontal_line(doc)
    footer_p = doc.add_paragraph()
    set_rtl(footer_p)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run('מסמך זה הוא תקציר של 15 מחקרים עצמאיים מעמיקים: ניתוח מתחרים, ארכיטקטורת סוכנים, מודלי AI, sandboxes, MCP, אבטחה, observability, תמחור, ועוד.')
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MUTED
    r.font.name = 'Calibri'

    out = '/home/user/landing-page-builder-skill/CodeForge-Overview.docx'
    doc.save(out)
    print(f'✅ Saved: {out}')


if __name__ == '__main__':
    build()
