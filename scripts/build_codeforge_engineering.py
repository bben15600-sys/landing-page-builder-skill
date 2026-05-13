#!/usr/bin/env python3
"""Engineering build plan for CodeForge — solo developer, internal tool, <$200/mo budget.

This is an engineering blueprint, not marketing. Sections:
1. Goals & non-goals
2. Functional requirements
3. Architecture (components, data flow, schema, API contracts)
4. Build plan (12 weeks, day-by-day for solo)
5. Costs (real numbers per service)
6. Open decisions
7. Risks & mitigations
8. Milestones
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x0B, 0x5E, 0xD7)
ACCENT = RGBColor(0xE8, 0x8B, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x72, 0x80)
SUCCESS = RGBColor(0x10, 0xA8, 0x5A)
DANGER = RGBColor(0xDC, 0x26, 0x26)
CODE_BG = "F3F4F6"
HEADER_BG = "1A1A2E"
ALT_ROW_BG = "F3F6FB"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def add_hr(doc, color='0B5ED7'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading(doc, text, level=1, rtl=True):
    sizes = {0: 28, 1: 20, 2: 15, 3: 12}
    size = sizes.get(level, 12)
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = PRIMARY if level <= 2 else DARK
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    if level == 1:
        add_hr(doc)
    return p


def para(doc, text, bold=False, color=None, size=11, italic=False, rtl=True):
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
    return p


def code_block(doc, text):
    """Monospace block, gray bg — for schemas, diagrams, code."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Consolas'
    run.font.color.rgb = DARK
    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if p.runs:
        p.runs[0].text = ''
    new_run = p.add_run(text)
    new_run.font.size = Pt(11)
    new_run.font.name = 'Calibri'
    return p


def make_table(doc, headers, rows, col_widths=None, header_bg=HEADER_BG):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tblPr = table._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_bg)
        cell_p = hdr_cells[i].paragraphs[0]
        set_rtl(cell_p)
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell_p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        if r_idx % 2 == 0:
            for c in cells:
                set_cell_bg(c, ALT_ROW_BG)
        for c_idx, val in enumerate(row):
            cell_p = cells[c_idx].paragraphs[0]
            set_rtl(cell_p)
            cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell_p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for r in table.rows:
                r.cells[i].width = Inches(w)
    return table


def callout(doc, title, body, bg="E7F1FE", border="0B5ED7", title_color=PRIMARY):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), border)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    p1 = cell.paragraphs[0]
    set_rtl(p1)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(title)
    r1.font.bold = True
    r1.font.size = Pt(11.5)
    r1.font.color.rgb = title_color
    r1.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Calibri'
    return table


def cover(doc):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, '1A1A2E')

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(80)
    r1 = p1.add_run('CodeForge')
    r1.font.size = Pt(56)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p2)
    r2 = p2.add_run('תכנית הנדסית לבנייה')
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0xE7, 0xF1, 0xFE)
    r2.font.name = 'Calibri'

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p3)
    p3.paragraph_format.space_before = Pt(30)
    r3 = p3.add_run('סולו דעוולופר · כלי פנימי · 12 שבועות · יעד <$150 לחודש')
    r3.font.size = Pt(12)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0xE7, 0xF1, 0xFE)
    r3.font.name = 'Calibri'

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p4)
    p4.paragraph_format.space_before = Pt(40)
    p4.paragraph_format.space_after = Pt(80)
    r4 = p4.add_run('מסמך הנדסי · אדריכלות · schemas · API contracts · עלויות · timeline')
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    r4.font.name = 'Calibri'


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    sectPr.append(bidi)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    cover(doc)
    doc.add_page_break()

    # ===== TOC =====
    heading(doc, 'תוכן עניינים', level=1)
    toc_items = [
        '1. מטרות המערכת ומה לא במטרות',
        '2. דרישות פונקציונליות וזרימות עבודה',
        '3. אדריכלות — רכיבים, schema, API contracts',
        '4. תכנית בנייה — 12 שבועות, יום אחר יום',
        '5. עלויות אמיתיות — לפי שירות ולפי שלב',
        '6. החלטות פתוחות — מה להחליט לפני התחלה',
        '7. סיכונים ואסטרטגיית מיתון',
        '8. Milestones, demos, exit criteria',
    ]
    for t in toc_items:
        bullet(doc, t)
    doc.add_page_break()

    # ===== 1. GOALS =====
    heading(doc, '1. מטרות המערכת', level=1)

    heading(doc, '1.1 מטרת על', level=2)
    para(doc,
         'לבנות לי סוכן AI אישי שמייצר ומתחזק קוד עבור הפרויקטים שלי, ללא תלות '
         'במוצרים מסחריים סגורים. הכלי הוא תשתית — לא מוצר. הוא מתפתח לפי הצרכים שלי, '
         'נשען על OSS ועל free tiers, ויודע להחליף את ספק ה-AI מאחורי הקלעים בלי שאני אצטרך לשנות זרימת עבודה.')

    heading(doc, '1.2 יעדים מדידים', level=2)
    make_table(doc,
        ['יעד', 'מדידה', 'יעד מספרי'],
        [
            ['Time-to-first-preview', 'מ-prompt בעברית עד אפליקציה רצה ב-iframe', '< 5 דקות'],
            ['עלות חודשית', 'סך כל השירותים + API', '< $150'],
            ['סוגי פרויקטים נתמכים', 'frameworks/runtimes שאני בונה איתם', '5+ (Next.js, Python, RN, scripts, docs)'],
            ['Vendor lock-in', 'מספר ספקי AI שהמערכת תומכת בהם', '≥ 3 (Anthropic, OpenAI, OSS)'],
            ['החזרה אחרי downtime', 'אם ספק יורד — להחליף ולהמשיך', '< 5 דקות'],
            ['Session retention', 'חזרה לפרויקט אחרי שבוע ולהמשיך לעבוד', '100%'],
            ['Self-dogfooding', 'אחוז זמן פיתוח שלי שמתבצע דרך הכלי', '≥ 60% תוך 3 חודשים'],
        ],
        col_widths=[2.2, 3.0, 1.8])

    heading(doc, '1.3 מה לא במטרות (Out of scope)', level=2)
    para(doc, 'ההחלטות האלו קריטיות — הן חוסכות חודשי עבודה ומגדירות מה לא לבנות:', italic=True, color=MUTED)
    out_of_scope = [
        'Multi-tenancy — אין משתמשים חוץ ממני, אין צורך בארגונים, צוותים, RBAC',
        'Billing/Stripe/subscriptions — אני משלם API ישירות לספקים',
        'SOC 2 / GDPR compliance — לא רלוונטי לכלי אישי',
        'ממשק עבור לא-מפתחים — אני מפתח, ה-UX יכול להיות technical',
        'Mobile native — Web-only, רספונסיבי מספיק',
        'Marketplace של templates / community / share — single user',
        'SSO / SAML / Enterprise auth — GitHub OAuth מסוים למשתמש אחד',
        'High-availability multi-region — single region, downtime סביר',
        'Real-time collaboration — אני לבד, אין co-editing',
        'Audit log חתום קריפטוגרפית — log רגיל ב-Postgres מספיק',
    ]
    for s in out_of_scope:
        bullet(doc, s)

    callout(doc,
        '⚠️ עיקרון מנחה',
        'כל בקשה שמגיעה תחת "out of scope" נדחית גם אם זה רעיון טוב. '
        'הסקופ קשיח עד MVP. אחרי MVP — דיון מחדש, לא לפני.',
        bg="FFF5E0", border="E88B00", title_color=ACCENT)

    doc.add_page_break()

    # ===== 2. FUNCTIONAL REQUIREMENTS =====
    heading(doc, '2. דרישות פונקציונליות', level=1)

    heading(doc, '2.1 User stories (פרסונה: אני)', level=2)
    stories = [
        ('US-1', 'פתיחת פרויקט חדש מ-prompt', 'אני כותב "build a Next.js todo app with Postgres" — תוך 5 דקות יש פרויקט עם git repo, README, שרת רץ, ו-URL ל-preview'),
        ('US-2', 'iterative editing בצ\'אט', 'אני אומר "הוסף dark mode" — המערכת מציגה diff, מאשרת איתי, מבצעת, ומראה preview מעודכן'),
        ('US-3', 'חזרה לפרויקט קיים', 'יום אחרי — פותח URL ל-session, רואה את כל ההיסטוריה, ממשיך מאיפה שעצרתי'),
        ('US-4', 'חיבור ל-GitHub שלי', 'מבקש "push to my GitHub as new repo" — נוצר repo פרטי, push, ו-Vercel deploy מקושר'),
        ('US-5', 'החלפת ספק AI מאמצע הצ\'אט', 'אני בוחר "use Gemini" — אותה שיחה ממשיכה עם Gemini, אותם tools, ההיסטוריה נשמרת'),
        ('US-6', 'הרצת bash בסביבה בטוחה', 'מבקש "run the tests" — הסוכן מריץ ב-sandbox, מחזיר output צבעוני, נכשל → מתקן → מריץ שוב'),
        ('US-7', 'בניית skill לשימוש חוזר', 'אומר "save this as a skill: deploy-to-vercel" — בפעם הבאה הסוכן זוכר את הצעדים'),
        ('US-8', 'עבודה offline על מודל מקומי', 'בלי אינטרנט — מתחבר ל-Qwen מקומי דרך LM Studio, ממשיך לעבוד'),
    ]
    make_table(doc,
        ['ID', 'כותרת', 'תיאור'],
        [[s[0], s[1], s[2]] for s in stories],
        col_widths=[0.7, 1.8, 4.5])

    heading(doc, '2.2 Core workflows', level=2)
    code_block(doc, """\
Flow A: Bootstrap a new project
─────────────────────────────────
  user prompt
    → LLM plans (file tree + commands)
    → user approves plan
    → tool: create_file × N
    → tool: bash (install deps)
    → tool: bash (run dev server)
    → WebContainer iframe → preview URL shown
    → save session to DB

Flow B: Iterative edit
─────────────────────────────────
  user message in existing session
    → LLM reads relevant files (tool: read_files)
    → LLM proposes edits as unified diff
    → user approves
    → tool: apply_edit
    → preview hot-reloads
    → optional: tool: run_tests

Flow C: Deploy
─────────────────────────────────
  user: "deploy this"
    → tool: github_create_repo (if not exists)
    → tool: github_push
    → tool: vercel_deploy
    → returns live URL

Flow D: Resume session
─────────────────────────────────
  open /s/{session_id}
    → load messages, files, tool_calls from DB
    → restore WebContainer from snapshot
    → ready to continue chat
""")

    heading(doc, '2.3 דרישות לא-פונקציונליות', level=2)
    make_table(doc,
        ['קטגוריה', 'דרישה', 'איך נמדד'],
        [
            ['Performance', 'First token < 2s, P95 streaming latency < 100ms', 'log + manual check'],
            ['Reliability', 'אם ספק AI נופל → fallback אוטומטי לספק חלופי', 'kill-test בסוף שלב 6'],
            ['Security', 'API keys לעולם לא ב-client. Sandbox בידוד מלא.', 'code review של key handling'],
            ['Maintainability', 'הוספת tool חדש = < 30 שורות קוד', 'דוגמה ב-tests'],
            ['Portability', 'self-host ב-VPS תוך פחות משעה', 'מסמך deploy + dry-run'],
            ['Observability', 'כל tool call נשמר עם duration + error', 'Postgres rows queryable'],
        ],
        col_widths=[1.5, 3.5, 2.0])

    doc.add_page_break()

    # ===== 3. ARCHITECTURE =====
    heading(doc, '3. אדריכלות', level=1)

    heading(doc, '3.1 דיאגרמת מערכת — high-level', level=2)
    code_block(doc, """\
┌─────────────────────────────────────────────────────────────────┐
│                    BROWSER (single user, me)                     │
│  ┌──────────┐  ┌────────────┐  ┌──────────┐  ┌──────────────┐   │
│  │ Chat UI  │  │ Monaco     │  │ File     │  │ WebContainer │   │
│  │ (msgs)   │  │ Editor     │  │ Tree     │  │ (iframe)     │   │
│  └────┬─────┘  └─────┬──────┘  └─────┬────┘  └──────┬───────┘   │
│       └──── all comm via React state ────┘             │         │
│                       │                                │         │
│                  SSE / WebSocket                       │         │
│                       │                          (postMessage)   │
└───────────────────────┼────────────────────────────────┼─────────┘
                        │                                │
                        ▼                                ▼
┌────────────────────────────────────────────┐  ┌──────────────────┐
│       Next.js 15 (Vercel Hobby — free)     │  │  StackBlitz      │
│  ┌──────────────────────────────────────┐  │  │  WebContainers   │
│  │  API Routes (Edge runtime where ok)  │  │  │  (free, OSS)     │
│  │  - /api/chat (SSE streaming)         │  │  │  Runs Node       │
│  │  - /api/sessions/...                 │  │  │  fully in browser│
│  │  - /api/tools/exec                   │  │  └──────────────────┘
│  │  - /api/github/* /api/vercel/*       │  │
│  └──────────────────────┬───────────────┘  │
│                         │                  │
│  ┌──────────────────────▼───────────────┐  │
│  │       Agent Loop (TypeScript)        │  │
│  │  ┌──────────┐  ┌────────────────┐    │  │
│  │  │ Provider │  │ Tool Executor  │    │  │
│  │  │ Router   │  │ (validates +   │    │  │
│  │  │          │  │  dispatches)   │    │  │
│  │  └─┬────────┘  └────┬───────────┘    │  │
│  └────┼────────────────┼────────────────┘  │
└───────┼────────────────┼───────────────────┘
        │                │
   ┌────▼──┐  ┌──────────▼────────┐  ┌──────────┐
   │ LLMs  │  │ Tool Adapters     │  │ Postgres │
   │       │  │ - GitHub API      │  │  (Neon   │
   │ - Cl  │  │ - Vercel API      │  │  free)   │
   │ - GPT │  │ - MCP servers     │  │          │
   │ - Gem │  │ - shell-in-sandbox│  │ sessions │
   │ - OSS │  │                   │  │ messages │
   └───────┘  └───────────────────┘  │ tool_log │
                                     └──────────┘
""")

    heading(doc, '3.2 רכיבים ואחריות', level=2)
    make_table(doc,
        ['רכיב', 'אחריות', 'מודול / קבצים'],
        [
            ['Web UI Shell', 'Layout, sidebar, routing, theme', 'app/(chat)/layout.tsx, components/shell/'],
            ['Chat UI', 'Message list, input, streaming tokens, tool-call UI', 'components/chat/'],
            ['Monaco wrapper', 'Code editing with file tree sync', 'components/editor/'],
            ['Preview pane', 'iframe to WebContainer, URL display', 'components/preview/'],
            ['Session API', 'CRUD על sessions, messages, files', 'app/api/sessions/'],
            ['Chat API (SSE)', 'streaming endpoint, agent loop entry', 'app/api/chat/route.ts'],
            ['Agent Loop', 'plan-act-observe, tool dispatch, halt condition', 'lib/agent/loop.ts'],
            ['Provider Router', 'Anthropic/OpenAI/Gemini adapters + routing', 'lib/llm/{anthropic,openai,gemini}.ts'],
            ['Tool Executor', 'validate args (zod), call adapter, format result', 'lib/tools/executor.ts'],
            ['Tool Adapters', 'read/write/bash/browser/github/vercel/mcp', 'lib/tools/adapters/*.ts'],
            ['Sandbox Bridge', 'postMessage protocol with WebContainer iframe', 'lib/sandbox/bridge.ts'],
            ['DB Layer', 'Drizzle queries, transactions', 'lib/db/{schema,queries}.ts'],
            ['Auth', 'GitHub OAuth (single-user allowlist)', 'lib/auth.ts (Auth.js)'],
        ],
        col_widths=[1.6, 3.0, 2.4])

    heading(doc, '3.3 Tech stack — בחירות עם נימוק', level=2)
    make_table(doc,
        ['שכבה', 'בחירה', 'למה דווקא זה'],
        [
            ['Framework', 'Next.js 15 (App Router)', 'SSR + API routes + Vercel = פריסה בלחיצה. Edge runtime לחלק מה-routes.'],
            ['Language', 'TypeScript strict', 'type safety מהותית כשעובדים עם LLM tool schemas'],
            ['Streaming/AI', 'Vercel AI SDK v5', 'streaming + typed tools + resumable streams מובנים'],
            ['UI', 'Tailwind + shadcn/ui', 'zero vendor lock, ownership של הקומפוננטות'],
            ['Editor', 'Monaco Editor', 'אותו editor של VS Code, חוויה מוכרת'],
            ['Sandbox', 'WebContainers (StackBlitz, free OSS)', 'רץ ב-browser, אפס עלות compute, אפס סיכון בידוד'],
            ['DB', 'Neon Postgres (free tier)', '3GB free, branching, serverless, אין downtime לפיתוח'],
            ['ORM', 'Drizzle', 'edge-compatible, type-safe, ללא overhead של Prisma'],
            ['Auth', 'Auth.js (NextAuth) + GitHub OAuth', 'free, allowlist של GitHub username שלך בלבד'],
            ['Hosting', 'Vercel Hobby', 'free עד 100GB bandwidth, מספיק לפרויקט אישי'],
            ['File storage', 'Cloudflare R2 (10GB free)', 'snapshots של sessions, אפס egress fees'],
            ['LLM ראשי', 'Anthropic Claude Sonnet 4.6', 'הכי טוב בקוד, prompt caching ב-90% הנחה'],
            ['LLM זול', 'Anthropic Haiku 4.5', 'tool calling פשוט, פי 10 זול'],
            ['LLM long-context', 'Gemini 3.1 (OpenRouter)', 'חלון 1M tokens, $1/M זול יחסית'],
            ['LLM OSS fallback', 'Qwen3-Coder via LM Studio', 'self-host על Mac, אפס עלות API'],
            ['Logging', 'Pino → Postgres', 'אין צורך ב-Sentry בשלב זה, log פשוט מספיק'],
            ['Jobs/Queue', '— (none initially)', 'הכל סנכרוני. אם נצטרך — Vercel Cron + DB queue'],
        ],
        col_widths=[1.2, 2.0, 3.8])

    heading(doc, '3.4 Database schema (Drizzle)', level=2)
    code_block(doc, """\
// lib/db/schema.ts
import { pgTable, text, timestamp, jsonb, integer, uuid, index } from 'drizzle-orm/pg-core';

export const sessions = pgTable('sessions', {
  id: uuid('id').primaryKey().defaultRandom(),
  title: text('title').notNull(),
  createdAt: timestamp('created_at').defaultNow().notNull(),
  updatedAt: timestamp('updated_at').defaultNow().notNull(),
  // user is implicit (single-user app), no user_id column
  archived: integer('archived').default(0).notNull(),
  // last provider/model used — for resume continuity
  lastProvider: text('last_provider'),
  lastModel: text('last_model'),
  // serialized WebContainer snapshot key (in R2)
  snapshotKey: text('snapshot_key'),
});

export const messages = pgTable('messages', {
  id: uuid('id').primaryKey().defaultRandom(),
  sessionId: uuid('session_id').references(() => sessions.id, { onDelete: 'cascade' }),
  role: text('role').notNull(),  // 'user' | 'assistant' | 'tool'
  content: jsonb('content').notNull(),  // array of content blocks
  createdAt: timestamp('created_at').defaultNow().notNull(),
  // tokens for cost accounting
  inputTokens: integer('input_tokens'),
  outputTokens: integer('output_tokens'),
  cachedTokens: integer('cached_tokens'),
  costCents: integer('cost_cents'),
}, (t) => ({
  sessionIdx: index('msg_session_idx').on(t.sessionId, t.createdAt),
}));

export const toolCalls = pgTable('tool_calls', {
  id: uuid('id').primaryKey().defaultRandom(),
  sessionId: uuid('session_id').references(() => sessions.id, { onDelete: 'cascade' }),
  messageId: uuid('message_id').references(() => messages.id),
  toolName: text('tool_name').notNull(),
  input: jsonb('input').notNull(),
  output: jsonb('output'),
  error: text('error'),
  durationMs: integer('duration_ms'),
  createdAt: timestamp('created_at').defaultNow().notNull(),
});

export const projects = pgTable('projects', {
  id: uuid('id').primaryKey().defaultRandom(),
  sessionId: uuid('session_id').references(() => sessions.id),
  name: text('name').notNull(),
  githubRepo: text('github_repo'),
  vercelProjectId: text('vercel_project_id'),
  liveUrl: text('live_url'),
  createdAt: timestamp('created_at').defaultNow().notNull(),
});

export const files = pgTable('files', {
  id: uuid('id').primaryKey().defaultRandom(),
  sessionId: uuid('session_id').references(() => sessions.id, { onDelete: 'cascade' }),
  path: text('path').notNull(),
  content: text('content').notNull(),
  updatedAt: timestamp('updated_at').defaultNow().notNull(),
}, (t) => ({
  pathIdx: index('files_session_path_idx').on(t.sessionId, t.path),
}));

export const costsDaily = pgTable('costs_daily', {
  day: text('day').primaryKey(),  // 'YYYY-MM-DD'
  provider: text('provider').notNull(),
  inputTokens: integer('input_tokens').notNull(),
  outputTokens: integer('output_tokens').notNull(),
  cachedTokens: integer('cached_tokens').notNull(),
  costCents: integer('cost_cents').notNull(),
});
""")

    heading(doc, '3.5 API contracts (internal)', level=2)
    code_block(doc, """\
POST /api/chat
  body: {
    sessionId: string,
    message: string,
    provider?: 'anthropic' | 'openai' | 'gemini' | 'local',
    model?: string,
  }
  response: text/event-stream
    data: { type: 'text-delta', delta: string }
    data: { type: 'tool-call-start', id: string, name: string, args: any }
    data: { type: 'tool-call-result', id: string, result: any }
    data: { type: 'usage', input: number, output: number, cached: number }
    data: { type: 'done', messageId: string }

POST /api/sessions
  body: { title?: string }
  response: { id: string, title: string, createdAt: string }

GET /api/sessions/:id
  response: { session, messages[], files[], project? }

POST /api/sessions/:id/snapshot
  body: { containerState: base64 }
  response: { snapshotKey: string }

POST /api/tools/exec   (internal — called by agent loop, not from browser)
  body: { sessionId: string, toolName: string, input: any }
  response: { output: any, durationMs: number }

GET /api/costs?from=YYYY-MM-DD&to=YYYY-MM-DD
  response: { totalCents: number, byProvider: {...}, dailyBreakdown: [...] }
""")

    heading(doc, '3.6 Sequence diagram — chat round-trip', level=2)
    code_block(doc, """\
User              UI               /api/chat        Anthropic        Tool Exec        WebContainer
  │                │                   │                │                 │                 │
  ├─message────────►                   │                │                 │                 │
  │                ├─POST(SSE)─────────►                │                 │                 │
  │                │                   ├─create stream──►                 │                 │
  │                │                   ◄──text delta────┤                 │                 │
  │                ◄──SSE: text-delta──┤                │                 │                 │
  │                │                   ◄──tool_use──────┤                 │                 │
  │                │                   ├─exec(toolName,args)──────────────►                 │
  │                │                   │                │                 ├─postMessage────►│
  │                │                   │                │                 ◄─────result──────┤
  │                │                   ◄────result──────┤                 │                 │
  │                ◄──SSE: tool-result─┤                │                 │                 │
  │                │                   ├─continue(stream w/ tool result)─►                 │
  │                │                   ◄────text delta──┤                 │                 │
  │                ◄──SSE: text-delta──┤                │                 │                 │
  │                │                   ◄────end_turn────┤                 │                 │
  │                │                   ├─save msg + tokens to DB          │                 │
  │                ◄──SSE: done────────┤                │                 │                 │
""")

    heading(doc, '3.7 Tool catalog (MVP set)', level=2)
    make_table(doc,
        ['Tool', 'מטרה', 'מבצע ב-'],
        [
            ['read_file', 'קריאת קובץ מ-session', 'DB + WebContainer FS'],
            ['write_file', 'יצירה/עדכון קובץ', 'DB + WebContainer FS'],
            ['list_files', 'tree של ה-project', 'WebContainer FS'],
            ['apply_diff', 'unified diff על קובץ קיים', 'WebContainer FS'],
            ['bash', 'הרצת פקודת shell', 'WebContainer'],
            ['http_fetch', 'GET/POST לכתובת', 'Server (Edge runtime)'],
            ['github_*', 'create_repo, push, list_repos', 'Octokit'],
            ['vercel_deploy', 'deploy מתוך WebContainer FS', 'Vercel API'],
            ['browser_screenshot', 'לוקח screenshot של preview', 'Playwright/Lambda (Phase 4+)'],
            ['save_skill', 'שומר שלבים לזיכרון', 'DB (skills table — Phase 8)'],
        ],
        col_widths=[1.8, 2.8, 2.4])

    doc.add_page_break()

    # ===== 4. BUILD PLAN =====
    heading(doc, '4. תכנית בנייה — 12 שבועות סולו', level=1)
    para(doc,
         'כל שבוע מציין deliverable מוחשי + קריטריון השלמה. הימים הם הנחות סבירות '
         'לסולו דעוולופר במשרה מלאה. אם זה side-project, הכפל ב-2.',
         italic=True, color=MUTED)

    weeks = [
        ('שבוע 1', 'Foundation',
         '• יום 1-2: GitHub repo + Next.js 15 scaffold + Tailwind + shadcn\n'
         '• יום 3: Neon DB + Drizzle schema + migrations רצות\n'
         '• יום 4: Auth.js עם GitHub OAuth + allowlist single user\n'
         '• יום 5: deploy ראשון ל-Vercel, env vars מוגדרים',
         'Localhost + Vercel preview עובדים, login דרך GitHub, יכול לכתוב ל-DB'),
        ('שבוע 2', 'Chat UI + Anthropic streaming',
         '• יום 6-7: Chat layout + message list + input\n'
         '• יום 8: Vercel AI SDK v5 + Anthropic adapter\n'
         '• יום 9: SSE endpoint /api/chat + streaming tokens\n'
         '• יום 10: messages persistence ל-Postgres',
         'יכול לקיים שיחה מלאה עם Claude, ההיסטוריה נשמרת ונטענת'),
        ('שבוע 3', 'Tool calls — first 3 tools',
         '• יום 11-12: zod schemas לכלים + dispatcher\n'
         '• יום 13: read_file + write_file (in-memory בשלב זה)\n'
         '• יום 14: bash דרך Node child_process (יוחלף בשבוע 4)\n'
         '• יום 15: tool_calls table + UI לתצוגת tool calls',
         'הסוכן יכול ליצור קובץ ולהריץ Node script, רואים את כל הקריאות ב-UI'),
        ('שבוע 4', 'WebContainers integration',
         '• יום 16-17: התקנת @webcontainer/api + iframe חבוי\n'
         '• יום 18: bridge postMessage לכלים\n'
         '• יום 19: file system sync (DB → WebContainer FS)\n'
         '• יום 20: dev server startup + preview URL',
         'אפליקציה שהסוכן יוצר רצה ב-WebContainer ומוצגת ב-iframe'),
        ('שבוע 5', 'Monaco editor + file tree',
         '• יום 21-22: Monaco wrapper + file tree component\n'
         '• יום 23: diff view (Monaco diff editor)\n'
         '• יום 24: עריכה ידנית מסונכרנת ל-WebContainer\n'
         '• יום 25: hot reload preview',
         'שלוש panes: chat | editor + tree | preview — כולן מסונכרנות'),
        ('שבוע 6', 'GitHub integration',
         '• יום 26-27: Octokit + create_repo tool\n'
         '• יום 28: clone repo INTO WebContainer\n'
         '• יום 29: git operations מתוך WebContainer (commit, push)\n'
         '• יום 30: project linking — session ↔ repo',
         'אפשר ליצור פרויקט, לעבוד עליו, ולעשות push לרפו פרטי שלך'),
        ('שבוע 7', 'Vercel deploy + project mgmt',
         '• יום 31-32: Vercel SDK + project creation\n'
         '• יום 33: deploy tool + URL ב-UI\n'
         '• יום 34: env vars sync (DB → Vercel)\n'
         '• יום 35: rollback / list deployments',
         'מ-prompt עד URL חי תוך < 10 דקות'),
        ('שבוע 8', 'Multi-provider router',
         '• יום 36-37: OpenAI adapter (Responses API)\n'
         '• יום 38: Gemini adapter (via OpenRouter — אחד API לכל ה-OSS)\n'
         '• יום 39: provider selector ב-UI + DB persistence per session\n'
         '• יום 40: tool schema normalization (Anthropic ↔ OpenAI tool format)',
         'אותה שיחה יכולה לעבור בין Claude/GPT/Gemini בלי לאבד context'),
        ('שבוע 9', 'Cost tracking + caching',
         '• יום 41: token accounting per message\n'
         '• יום 42: prompt caching עם Anthropic (5min + 1h)\n'
         '• יום 43: costs_daily table + /api/costs endpoint\n'
         '• יום 44-45: costs dashboard ב-UI עם אזעקות לפני $X',
         'רואה בדיוק כמה הוצאתי החודש, ניתן לקבוע cap'),
        ('שבוע 10', 'MCP client (basic)',
         '• יום 46-47: MCP stdio + SSE client\n'
         '• יום 48: רישום שרתים ב-config + UI להפעלה/כיבוי\n'
         '• יום 49: חיבור 2-3 servers (filesystem, github official, sentry)\n'
         '• יום 50: עטיפת MCP tools ב-tool executor',
         'יכול להפעיל כלי MCP בתוך שיחה, hash-pinning בסיסי'),
        ('שבוע 11', 'Memory + skills',
         '• יום 51-52: skills table + save_skill tool\n'
         '• יום 53: auto-summarization של שיחות ארוכות (compaction)\n'
         '• יום 54: project context file (CodeForge.md per session)\n'
         '• יום 55: UI לעריכת skills ידנית',
         'הסוכן מתחיל להצטבר ידע, ניתן לעבד skills'),
        ('שבוע 12', 'Polish + dogfooding',
         '• יום 56-57: keyboard shortcuts, dark mode, mobile breakpoints\n'
         '• יום 58: bug bash — תקן את 5 הדברים הכי מעצבנים\n'
         '• יום 59: deploy לדומיין פרטי, snapshots של sessions ל-R2\n'
         '• יום 60: סגירת MVP. שבוע אחד dogfooding ללא הוספות.',
         'MVP חי, אני משתמש בו לפיתוח אמיתי במקום Cursor/Claude Code'),
    ]
    make_table(doc,
        ['שבוע', 'נושא', 'משימות יומיות', 'קריטריון השלמה'],
        [[w[0], w[1], w[2], w[3]] for w in weeks],
        col_widths=[0.7, 1.4, 3.5, 1.8])

    callout(doc,
        '🛑 Discipline check',
        'כל שבוע — עוצר ב-end-of-week ושואל: האם ה-deliverable עומד? אם לא — '
        'דוחה את כל מה שמתחתיו ולא מוסיף scope חדש. סלייד = פיצ\'ר שהורד, '
        'לא דדליין שזז.',
        bg="FFF5E0", border="E88B00", title_color=ACCENT)

    doc.add_page_break()

    # ===== 5. COSTS =====
    heading(doc, '5. עלויות אמיתיות', level=1)

    heading(doc, '5.1 חד-פעמי (setup)', level=2)
    make_table(doc,
        ['פריט', 'עלות', 'הערות'],
        [
            ['דומיין (אופציונלי)', '$10-12/שנה', 'codeforge.dev / .app'],
            ['GitHub account', '$0', 'free tier מספיק'],
            ['Vercel account', '$0', 'Hobby plan'],
            ['Neon account', '$0', 'free tier'],
            ['Cloudflare account', '$0', 'R2 free tier'],
            ['Anthropic API account', '$0 (pay-as-you-go)', 'מינ\' deposit $5'],
            ['OpenAI API account', '$0 (pay-as-you-go)', 'מינ\' deposit $5'],
            ['OpenRouter (אגרגטור)', '$0 (pay-as-you-go)', 'גישה ל-Gemini/Qwen בלי לפתוח 5 חשבונות'],
            ['סה"כ setup', '$0-12', 'בלי דומיין: אפס'],
        ],
        col_widths=[2.4, 2.0, 3.0])

    heading(doc, '5.2 חודשי בפיתוח (שבועות 1-12)', level=2)
    para(doc, 'הנחה: שימוש מתון תוך כדי פיתוח — ~20 sessions/חודש, ~30K tokens IO ממוצע.',
         italic=True, color=MUTED)
    make_table(doc,
        ['שירות', 'תכנית', 'כלול', 'עלות חודשית'],
        [
            ['Vercel', 'Hobby', '100GB bandwidth, builds', '$0'],
            ['Neon Postgres', 'Free', '3GB storage, 100h compute', '$0'],
            ['Cloudflare R2', 'Free', '10GB storage, 1M ops', '$0'],
            ['GitHub', 'Free', 'unlimited private repos', '$0'],
            ['Anthropic Claude', 'pay-as-you-go', 'Sonnet 4.6 + Haiku', '$25-50'],
            ['OpenAI', 'pay-as-you-go', 'GPT-5.5 לבדיקות', '$5-10'],
            ['OpenRouter', 'pay-as-you-go', 'Gemini + OSS לבדיקות', '$3-5'],
            ['Domain', 'annual / 12', '— (אם רכשת)', '$1'],
            ['סה"כ', '', '', '$34-66'],
        ],
        col_widths=[1.6, 1.4, 2.5, 1.3])

    heading(doc, '5.3 חודשי אחרי MVP (שימוש כבד יומיומי)', level=2)
    para(doc, 'הנחה: dogfooding מלא, ~50 sessions/חודש, ~80K tokens IO ממוצע, prompt caching מופעל.',
         italic=True, color=MUTED)
    make_table(doc,
        ['שירות', 'עלות חודשית', 'תרחיש להגדלה'],
        [
            ['Vercel Hobby → Pro?', '$0 → $20', 'רק אם תעבור bandwidth limit'],
            ['Neon Free → Launch', '$0 → $19', 'רק אם DB עולה על 3GB'],
            ['Anthropic (Sonnet primary)', '$60-90', 'תלוי בשימוש; caching חוסך 70%'],
            ['Anthropic (Haiku for edits)', '$3-7', 'משימות פשוטות'],
            ['Anthropic (Opus for planning)', '$10-20', 'שימוש סלקטיבי'],
            ['OpenRouter (fallback/exper)', '$5-10', 'Gemini long-context'],
            ['R2 storage', '$0-3', 'snapshots של sessions'],
            ['סה"כ ריאלי', '$78-149', 'מתחת ל-$200/חודש הנחיה'],
        ],
        col_widths=[2.3, 1.7, 3.0])

    heading(doc, '5.4 פירוט עלות API — איך החישוב נעשה', level=2)
    code_block(doc, """\
Sonnet 4.6 pricing (May 2026):
  Input:           $3.00 / 1M tokens
  Output:          $15.00 / 1M tokens
  Cache write:     $3.75 / 1M tokens (1.25× input)
  Cache read:      $0.30 / 1M tokens (0.1× input — 90% discount)

Example session (typical):
  System prompt:    10K tokens (cached after first turn)
  Project files:    15K tokens (cached after read)
  User msg:         500 tokens
  Tool results:     5K tokens
  LLM output:       2K tokens
  Total per turn:   ~32K tokens

First turn (no cache):
  Input cost:  30K × $3/M  = $0.090
  Output cost: 2K × $15/M  = $0.030
  Turn cost:                = $0.120

Cached turn (90% input is cache hit):
  Cache read:  27K × $0.30/M = $0.008
  Fresh input: 3K × $3/M     = $0.009
  Output:      2K × $15/M    = $0.030
  Turn cost:                  = $0.047

Session of 10 turns:
  First turn:  $0.120
  9 × cached:  $0.420
  Total:       $0.540

Monthly (50 sessions): ~$27 — well within budget.

Cost-cutting tricks the system enforces:
  1. Sonnet for code, Haiku for ack/parse (10× cheaper per token)
  2. Auto-cache system prompt + project context (90% off)
  3. Compact conversations > 100 messages into summary
  4. Route to Gemini for >200K token contexts (Sonnet starts to degrade + cost)
  5. Hard monthly cap at $150 — agent stops, user must lift cap
""")

    heading(doc, '5.5 איך להוריד עוד', level=2)
    bullet(doc, 'Qwen3-Coder 32B מקומי על Mac M-series דרך LM Studio: ~$0 API. איטי יותר אבל חינם.')
    bullet(doc, 'Groq לאינפרנס מהיר ב-$0.50/M tokens (Llama / Qwen) — פי 6 זול מ-Sonnet.')
    bullet(doc, 'OpenRouter עם DeepSeek-V3.2 — $0.27/M input, איכות קרובה ל-Sonnet במשימות פשוטות.')
    bullet(doc, 'Prompt caching לא רק לקלוד — Gemini 2.5+ ו-DeepSeek גם תומכים, הנחות 75-90%.')
    bullet(doc, 'Batch API (Anthropic, OpenAI): 50% הנחה למשימות לא-זמן-אמת (analysis, summarization).')

    callout(doc,
        '💰 שורה תחתונה על עלויות',
        'התקציב היעד <$200/חודש בר-השגה ברמה ניכרת. במהלך פיתוח: $34-66. '
        'אחרי MVP בשימוש יומיומי: $78-149. אם תעבור — סימן ל-A) להפעיל caching יותר אגרסיבי, '
        'B) לעבור ל-Haiku לרוב המשימות, C) להפעיל local model לעבודה rutínit.',
        bg="E7FBE7", border="10A85A", title_color=SUCCESS)

    doc.add_page_break()

    # ===== 6. OPEN DECISIONS =====
    heading(doc, '6. החלטות פתוחות לפני התחלה', level=1)
    para(doc, 'לכל החלטה: ההמלצה הראשונית + הסיבה. אפשר לחרוג, אבל כדאי לחרוג מודע.',
         italic=True, color=MUTED)

    make_table(doc,
        ['#', 'החלטה', 'אופציות', 'המלצה', 'נימוק'],
        [
            ['D1', 'Auth provider', 'Auth.js / Clerk / Better-Auth / custom', 'Auth.js', 'חינם, GitHub OAuth מובנה, single-user מספיק'],
            ['D2', 'Sandbox runtime', 'WebContainers / E2B / Daytona / local Docker', 'WebContainers', 'אפס עלות compute, רץ בדפדפן'],
            ['D3', 'LLM ראשי', 'Claude Sonnet / GPT / Gemini', 'Claude Sonnet 4.6', 'הכי טוב בקוד, prompt caching טוב'],
            ['D4', 'Streaming library', 'Vercel AI SDK / LangChain / custom', 'Vercel AI SDK v5', 'pure ergonomics, תמיכה ב-tools משובחת'],
            ['D5', 'UI components', 'shadcn / MUI / Mantine / Park UI', 'shadcn/ui', 'אתה הבעלים של הקוד, Tailwind native'],
            ['D6', 'Editor', 'Monaco / CodeMirror / Lexical', 'Monaco', 'אותו של VS Code, מוכר ל-LLMs'],
            ['D7', 'ORM', 'Drizzle / Prisma / Kysely', 'Drizzle', 'Edge-compatible, type-safe, fast'],
            ['D8', 'State (client)', 'Zustand / Jotai / Redux / React state', 'Zustand', 'מינימלי, ללא boilerplate'],
            ['D9', 'תמיכה ב-OSS models', 'OpenRouter / Together / Groq / local', 'OpenRouter + LM Studio', 'OpenRouter ל-cloud, LM Studio ל-local'],
            ['D10', 'Logging', 'Sentry / Axiom / Pino-to-DB / nothing', 'Pino → Postgres', 'בשלב זה אין צורך בכלי חיצוני'],
            ['D11', 'Testing', 'Vitest + Playwright / Jest / nothing', 'Vitest + Playwright', 'Vitest לליבה, Playwright ל-E2E של flows'],
            ['D12', 'Monorepo?', 'Turborepo / single Next.js / pnpm workspaces', 'single Next.js', 'אל תפזר עד שתצטרך'],
        ],
        col_widths=[0.4, 1.5, 2.0, 1.2, 2.4])

    doc.add_page_break()

    # ===== 7. RISKS =====
    heading(doc, '7. סיכונים ואסטרטגיית מיתון', level=1)
    make_table(doc,
        ['#', 'סיכון', 'סבירות', 'השפעה', 'מיתון'],
        [
            ['R1', 'עלויות API מתפוצצות בלי לשים לב', 'בינוני', 'גבוה', 'Hard cap בקוד, costs dashboard בכל login, alerts ב-$50/100/150'],
            ['R2', 'WebContainers לא תומך ב-X חיוני', 'נמוך', 'בינוני', 'Fallback: Node child_process מקומי + E2B SDK ($) במידת הצורך'],
            ['R3', 'Scope creep — להוסיף Stripe/multi-tenant', 'גבוה', 'גבוה', 'Out-of-scope list קשיח, כל בקשה לפיצ\'ר נכתבת ב-backlog.md לא בקוד'],
            ['R4', 'איבוד עניין באמצע (פרויקט סולו)', 'גבוה', 'גבוה', 'פיצול ל-2-week sprints עם demo מוחשי בסוף כל אחד; שיתוף עם חבר לחיצים'],
            ['R5', 'Breaking change ב-API של ספק AI', 'נמוך', 'בינוני', 'Provider adapter layer מהיום הראשון; pin SDK versions'],
            ['R6', 'Sonnet/Opus עוברים deprecation', 'נמוך', 'בינוני', 'Multi-provider מבטיח החלפה ב-2 שעות'],
            ['R7', 'נעילת vendor של Vercel', 'נמוך', 'נמוך', 'הקוד הוא Next.js תקני, deploy ל-Cloudflare/self-host אפשרי'],
            ['R8', 'אובדן sessions בגלל crash של Neon free', 'נמוך מאוד', 'בינוני', 'pg_dump שבועי ל-R2 כסקריפט cron'],
            ['R9', 'Tool injection מ-LLM (כתב bash הרסני)', 'בינוני', 'גבוה', 'allowlist ל-bash, deny rm -rf / sudo / חיבור חיצוני, אישור לפני ביצוע'],
            ['R10', 'בנייה ארוכה מדי ויציאה ל-shelf', 'גבוה', 'גבוה', 'ה-MVP חייב לרוץ עד שבוע 12. אחרי זה משתמשים בו 30 יום לפני להוסיף.'],
        ],
        col_widths=[0.4, 2.6, 0.9, 0.9, 2.7])

    doc.add_page_break()

    # ===== 8. MILESTONES =====
    heading(doc, '8. Milestones, demos, exit criteria', level=1)
    make_table(doc,
        ['Milestone', 'שבוע', 'Demo', 'exit criteria'],
        [
            ['M1: Hello agent', '2', 'שיחה רצה עם Claude ב-localhost', 'יכול לחזור לשיחה אחרי refresh'],
            ['M2: First tool', '3', 'הסוכן יוצר קובץ ומריץ Node', 'tool_calls נשמרים, UI מציג'],
            ['M3: Live preview', '4', 'הסוכן בונה todo app, רץ ב-iframe', 'preview hot-reload ב-< 2s'],
            ['M4: Full IDE feel', '5', 'editor + preview + chat מסונכרנים', 'עריכה ידנית משתקפת ב-preview'],
            ['M5: GitHub flow', '6', 'session → repo פרטי + clone חזרה', 'יכול לחדש עבודה ב-session ישן'],
            ['M6: Deploy', '7', 'מ-prompt עד URL חי', 'rollback עובד'],
            ['M7: Multi-LLM', '8', 'אותה שיחה עוברת בין Claude/GPT/Gemini', 'tool schemas מתורגמים אוטומטית'],
            ['M8: Cost-aware', '9', 'dashboard עם רוב הוצאה והתראות', 'caching מופעל, חיסכון נמדד'],
            ['M9: MCP-ready', '10', 'GitHub MCP server רץ בתוך שיחה', 'hash-pinning + allowlist'],
            ['M10: Memory', '11', 'project context + skill saved/reused', 'compaction עובד על שיחה ארוכה'],
            ['MVP DONE', '12', 'אני משתמש בכלי במקום Claude Code/Cursor', '30 ימי dogfooding ללא הוספות'],
        ],
        col_widths=[1.6, 0.7, 2.5, 2.7])

    heading(doc, '8.1 Definition of MVP done', level=2)
    para(doc, 'ה-MVP מוכן כש:')
    bullet(doc, 'אני יכול לבנות todo app מאפס תוך פחות מ-10 דקות')
    bullet(doc, 'אני יכול לחדש לעבוד על פרויקט מלפני שבוע ולהמשיך אותו')
    bullet(doc, 'אני יכול להעביר שיחה בין Claude ל-Gemini בלי לאבד context')
    bullet(doc, 'יש לי dashboard עלויות עדכני')
    bullet(doc, 'אני עומד מתחת ל-$150/חודש בשימוש יומי')
    bullet(doc, 'במהלך 30 ימי dogfooding לא הייתי צריך לחזור ל-Cursor/Claude Code')

    callout(doc,
        '🏁 השלב הבא — אחרי MVP',
        'ה-30 ימי dogfooding הם הקריטיים. הם אומרים לך מה באמת חסר — לא מה שחשבת '
        'בשלב התכנון. אחרי 30 יום: כתוב backlog מפורט מבוסס שימוש אמיתי, ותכנן '
        'Phase 2 (מה שבמסמך הזה מצוין כ-out-of-scope — קצת ממנו עשוי לעלות לפה).',
        bg="E7F1FE", border="0B5ED7", title_color=PRIMARY)

    add_hr(doc, color='6B7280')
    footer_p = doc.add_paragraph()
    set_rtl(footer_p)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run('מסמך הנדסי · גרסה 1.0 · לסולו דעוולופר · כלי פנימי · יעד <$150/חודש')
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = MUTED
    r.font.name = 'Calibri'

    out = '/home/user/landing-page-builder-skill/CodeForge-Build-Plan.docx'
    doc.save(out)
    print(f'✅ Saved: {out}')


if __name__ == '__main__':
    build()
